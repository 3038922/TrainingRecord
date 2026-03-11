# -*- coding: utf-8 -*-
"""
根据《2025全年训练日期汇总.xlsx》中的训练日期，
结合《机器人社团训练记录表模板.docx》，
按月份生成训练记录 Word 文件。

当前默认仅生成 2 月，用于调试。
确认无误后，把 ONLY_TEST_FEBRUARY = False，即可生成 1~12 月全部文件。
"""

from pathlib import Path
from copy import deepcopy
import math
import re
from datetime import datetime, date

import openpyxl
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn


# =========================
# 可根据实际情况修改的配置
# =========================

EXCEL_FILE = "2025全年训练日期汇总.xlsx"
TEMPLATE_FILE = "机器人社团训练记录表模板.docx"
OUTPUT_DIR = "训练记录输出"

# True = 只生成2月，便于调试
# False = 生成1月~12月
ONLY_TEST_FEBRUARY = False

MONTH_SHEET_NAMES = {
    1: "1月",
    2: "2月",
    3: "3月",
    4: "4月",
    5: "5月",
    6: "6月",
    7: "7月",
    8: "8月",
    9: "9月",
    10: "10月",
    11: "11月",
    12: "12月",
}

# 你确认后的训练时间映射
TIME_RANGE_MAP = {
    0.5: "17:30 - 18:00",
    1:   "17:00 - 18:00",
    3.5: "17:30 - 21:00",
    4:   "17:00 - 21:00",
    8:   "13:00 - 21:00",
}

WEEKDAY_MAP = {
    0: "星期一",
    1: "星期二",
    2: "星期三",
    3: "星期四",
    4: "星期五",
    5: "星期六",
    6: "星期日",
}


# =========================
# 工具函数
# =========================

def normalize_space(text: str) -> str:
    if text is None:
        return ""
    return re.sub(r"\s+", "", str(text)).strip()


def safe_float(value, default=None):
    if value is None or value == "":
        return default
    try:
        return float(value)
    except Exception:
        return default


def parse_excel_date(value):
    """
    兼容 Excel 日期、date/datetime、字符串日期。
    """
    if value is None or value == "":
        return None

    if isinstance(value, datetime):
        return value

    if isinstance(value, date):
        return datetime(value.year, value.month, value.day)

    if isinstance(value, (int, float)):
        try:
            from openpyxl.utils.datetime import from_excel
            dt = from_excel(value)
            if isinstance(dt, datetime):
                return dt
            if isinstance(dt, date):
                return datetime(dt.year, dt.month, dt.day)
        except Exception:
            pass

    if isinstance(value, str):
        text = value.strip()
        if not text:
            return None

        formats = [
            "%Y-%m-%d",
            "%Y/%m/%d",
            "%Y.%m.%d",
            "%Y-%m-%d %H:%M:%S",
            "%Y/%m/%d %H:%M:%S",
            "%Y.%m.%d %H:%M:%S",
        ]

        for fmt in formats:
            try:
                return datetime.strptime(text, fmt)
            except Exception:
                pass

        text2 = text.replace("年", "-").replace("月", "-").replace("日", "")
        for fmt in ["%Y-%m-%d", "%Y-%m-%d %H:%M:%S"]:
            try:
                return datetime.strptime(text2, fmt)
            except Exception:
                pass

    return None


def format_date_parts(dt):
    year = dt.year
    month = dt.month
    day = dt.day
    weekday = WEEKDAY_MAP[dt.weekday()]
    return year, month, day, weekday


def get_time_range_from_duration(duration):
    if duration is None:
        return "       -       "

    duration = round(float(duration), 2)

    if duration in TIME_RANGE_MAP:
        return TIME_RANGE_MAP[duration]

    for k in TIME_RANGE_MAP:
        if math.isclose(duration, k, abs_tol=0.01):
            return TIME_RANGE_MAP[k]

    hours = int(duration)
    minutes = int(round((duration - hours) * 60))
    return f"约{hours}小时{minutes}分钟"


def clear_paragraph_keep_format(paragraph):
    """
    清空段落文字，但尽量保留段落级格式。
    """
    if paragraph.runs:
        for run in paragraph.runs:
            run.text = ""
    else:
        paragraph.add_run("")


def set_paragraph_text_with_font(paragraph, text, font_name=None, font_size_pt=None, bold=None):
    """
    清空段落后写入新文字，并可指定字体/字号。
    """
    clear_paragraph_keep_format(paragraph)
    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run("")
    run.text = text

    if font_name:
        run.font.name = font_name
        r = run._element.rPr
        if r is None:
            r = run._element.get_or_add_rPr()
        r.rFonts.set(qn("w:eastAsia"), font_name)

    if font_size_pt:
        run.font.size = Pt(font_size_pt)

    if bold is not None:
        run.bold = bold


def set_cell_paragraph_text(cell, text, font_name=None, font_size_pt=None, bold=None):
    """
    修改单元格首段内容，尽量保留单元格、段落本身结构。
    不用 cell.text，避免把格式全部打掉。
    """
    if cell.paragraphs:
        p = cell.paragraphs[0]
    else:
        p = cell.add_paragraph()

    set_paragraph_text_with_font(
        p,
        text,
        font_name=font_name,
        font_size_pt=font_size_pt,
        bold=bold
    )


def load_month_records_from_sheet(ws):
    headers = {}
    for col in range(1, ws.max_column + 1):
        value = ws.cell(row=1, column=col).value
        if value is not None:
            headers[normalize_space(value)] = col

    date_col = None
    duration_col = None

    for header_text, col_idx in headers.items():
        if ("训练日期" in header_text) or (header_text == "日期"):
            date_col = col_idx

        if ("训练时长" in header_text) or ("时长" in header_text):
            duration_col = col_idx

    if date_col is None:
        raise ValueError(f"工作表 {ws.title} 未找到日期列，表头有：{list(headers.keys())}")
    if duration_col is None:
        raise ValueError(f"工作表 {ws.title} 未找到训练时长列，表头有：{list(headers.keys())}")

    records = []

    for row in range(2, ws.max_row + 1):
        raw_date = ws.cell(row=row, column=date_col).value
        raw_duration = ws.cell(row=row, column=duration_col).value

        dt = parse_excel_date(raw_date)
        if dt is None:
            continue

        duration = safe_float(raw_duration, default=None)

        records.append({
            "date": dt,
            "duration": duration,
        })

    return records


def find_template_blocks(template_doc):
    """
    找出模板中用于重复复制的一组元素：
    - 标题段落
    - 标题后的空白段落
    - 主表格
    - 模板自带分页符段落（如果有）
    不复制最后那个纯空白尾段，避免多余空白页。
    """
    body_children = list(template_doc._element.body)

    # 去掉 sectPr
    content_children = [x for x in body_children if not x.tag.endswith("sectPr")]

    block_title = None
    block_blank = None
    block_table = None
    block_page_break = None

    for child in content_children:
        tag = child.tag.split("}")[-1]
        if tag == "p" and block_title is None:
            block_title = child
            continue

        if tag == "p" and block_title is not None and block_blank is None:
            block_blank = child
            continue

        if tag == "tbl" and block_table is None:
            block_table = child
            continue

        if tag == "p" and block_table is not None:
            # 找模板自带分页符
            xml = child.xml
            if 'w:br w:type="page"' in xml:
                block_page_break = child
                break

    if block_title is None or block_table is None:
        raise ValueError("模板结构识别失败：未找到标题或主表格。")

    return block_title, block_blank, block_table, block_page_break


def remove_all_body_content_except_sectPr(doc):
    """
    清空正文内容，只保留 sectPr。
    """
    body = doc._element.body
    children = list(body)
    for child in children:
        if child.tag.endswith("sectPr"):
            continue
        body.remove(child)


def append_block_copy(doc, block_elems):
    """
    向文档末尾追加一组模板元素副本。
    """
    body = doc._element.body
    sectPr = None

    for child in list(body):
        if child.tag.endswith("sectPr"):
            sectPr = child
            break

    for elem in block_elems:
        if elem is not None:
            if sectPr is not None:
                body.insert(body.index(sectPr), deepcopy(elem))
            else:
                body.append(deepcopy(elem))


def is_target_record_table(table):
    """
    判断是否是训练记录主表。
    """
    if len(table.rows) == 0:
        return False

    row0_text = "".join(cell.text for cell in table.rows[0].cells)
    row0_norm = normalize_space(row0_text)
    return ("日期" in row0_norm) and ("星期" in row0_norm) and ("训练时段" in row0_norm)


def force_title_style(paragraph):
    """
    强制标题字体为：方正简宋小标，二号。
    二号约 22pt。
    """
    set_paragraph_text_with_font(
        paragraph,
        "机器人工作室训练记录表",
        font_name="方正简宋小标",
        font_size_pt=22,
        bold=False
    )


def fill_table_first_row(table, dt, duration):
    """
    按模板第一行单元格填写：
    cell[0] 日期
    cell[2] 星期
    cell[4] 训练时段
    """
    year, month, day, weekday = format_date_parts(dt)
    time_range = get_time_range_from_duration(duration)

    row = table.rows[0]

    set_cell_paragraph_text(row.cells[0], f"日期：{year}年{month}月{day}日")
    set_cell_paragraph_text(row.cells[2], weekday)
    set_cell_paragraph_text(row.cells[4], f"训练时段：{time_range}")


def fill_generated_doc_records(doc, records):
    """
    给所有页面填值，并修正标题字体。
    """
    target_tables = [t for t in doc.tables if is_target_record_table(t)]
    print(f"[调试] 文档中匹配到 {len(target_tables)} 个训练记录表格")

    if len(target_tables) < len(records):
        raise ValueError(
            f"文档中找到的训练记录表格数({len(target_tables)})小于记录数({len(records)})，"
            f"说明模板复制或表格识别可能有问题。"
        )

    # 标题段落一般就是所有非空正文段落里的这些标题
    title_paragraphs = [p for p in doc.paragraphs if normalize_space(p.text) == "机器人工作室训练记录表"]

    if len(title_paragraphs) < len(records):
        print(f"[警告] 识别到的标题数 {len(title_paragraphs)} 少于记录数 {len(records)}，但不影响主表填写。")

    for i, rec in enumerate(records):
        if i < len(title_paragraphs):
            force_title_style(title_paragraphs[i])

        fill_table_first_row(target_tables[i], rec["date"], rec["duration"])


def build_month_doc(month_num, records, template_path, output_dir):
    if not records:
        print(f"[跳过] {month_num}月没有有效训练记录。")
        return

    # 以模板本身为基础，避免样式丢失
    doc = Document(template_path)

    # 识别模板中需要重复的内容块
    block_title, block_blank, block_table, block_page_break = find_template_blocks(doc)

    # 清空正文，只保留节设置
    remove_all_body_content_except_sectPr(doc)

    # 逐条记录构建页面
    for idx, _ in enumerate(records):
        elems = [block_title, block_blank, block_table]

        # 不是最后一条时，才追加模板自带分页符
        if idx != len(records) - 1 and block_page_break is not None:
            elems.append(block_page_break)

        append_block_copy(doc, elems)

    # 填写具体内容
    fill_generated_doc_records(doc, records)

    output_dir.mkdir(parents=True, exist_ok=True)
    output_file = output_dir / f"机器人社团训练记录表_{month_num}月.docx"
    doc.save(output_file)
    print(f"[完成] 已生成：{output_file}")


def main():
    base_dir = Path(".")
    excel_path = base_dir / EXCEL_FILE
    template_path = base_dir / TEMPLATE_FILE
    output_dir = base_dir / OUTPUT_DIR

    if not excel_path.exists():
        raise FileNotFoundError(f"未找到 Excel 文件：{excel_path}")
    if not template_path.exists():
        raise FileNotFoundError(f"未找到模板文件：{template_path}")

    wb = openpyxl.load_workbook(excel_path, data_only=True)

    months = [2] if ONLY_TEST_FEBRUARY else list(range(1, 13))

    for month in months:
        sheet_name = MONTH_SHEET_NAMES[month]
        if sheet_name not in wb.sheetnames:
            print(f"[跳过] Excel 中不存在工作表：{sheet_name}")
            continue

        ws = wb[sheet_name]
        records = load_month_records_from_sheet(ws)
        records = [r for r in records if r["date"].month == month]

        print(f"[调试] {month}月读取到 {len(records)} 条记录")

        if not records:
            print(f"[跳过] {month}月没有可用记录。")
            continue

        build_month_doc(month, records, template_path, output_dir)


if __name__ == "__main__":
    main()